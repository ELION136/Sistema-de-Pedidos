from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse
from reportlab.lib.pagesizes import A4, landscape
from django.contrib.auth.decorators import login_required
from django.db.models import Sum, Count, Q
from django.core.paginator import Paginator
import pandas as pd
from decimal import Decimal

from .models import PedidoPersonalizado, ItemPedidoPersonalizado
from .forms import (
    PedidoPersonalizadoForm, ItemPedidoPersonalizadoForm,
    FiltrosPedidoPersonalizadoForm, ImportarExcelPersonalizadoForm
)
from apps.clientes.models import Cliente

# ── 1. NUEVA dashboard_personalizados ────────────────────────────────────────
@login_required
def dashboard_personalizados(request):
    # ── Filtros del resumen ───────────────────────────────────────────────
    tipo_prenda  = request.GET.get('tipo_prenda', '').strip()
    talla        = request.GET.get('talla', '').strip()
    genero       = request.GET.get('genero', '').strip()
    beneficiario = request.GET.get('beneficiario', '').strip()
    cliente_txt  = request.GET.get('cliente', '').strip()
 
    # ── Stats globales (siempre sobre todos los pedidos) ──────────────────
    todos          = PedidoPersonalizado.objects.all()
    total_personas = todos.count()
    total_prendas  = ItemPedidoPersonalizado.objects.aggregate(t=Sum('cantidad'))['t'] or 0
    total_aportes  = todos.aggregate(t=Sum('aporte'))['t'] or 0
    total_saldos   = todos.aggregate(t=Sum('saldo'))['t'] or 0
 
    # ── Resumen de prendas enriquecido con beneficiario y cliente ─────────
    # Anotamos directamente sobre ItemPedidoPersonalizado para incluir
    # los datos del pedido (beneficiario, cliente, categoría, gestión)
    items_qs = (
        ItemPedidoPersonalizado.objects
        .select_related('pedido', 'pedido__cliente')
        .values(
            'tipo_prenda',
            'talla',
            'genero',
            'pedido__nombre_completo',   # beneficiario
            'pedido__cliente__nombre',   # institución
            'pedido__categoria',
            'pedido__gestion',
        )
        .annotate(total=Sum('cantidad'))
        .order_by('tipo_prenda', 'talla', 'genero', 'pedido__nombre_completo')
    )
 
    # Aplicar filtros opcionales
    if tipo_prenda:
        items_qs = items_qs.filter(tipo_prenda=tipo_prenda)
    if talla:
        items_qs = items_qs.filter(talla=talla)
    if genero:
        items_qs = items_qs.filter(genero=genero)
    if beneficiario:
        items_qs = items_qs.filter(pedido__nombre_completo__icontains=beneficiario)
    if cliente_txt:
        items_qs = items_qs.filter(pedido__cliente__nombre__icontains=cliente_txt)
 
    # Normalizar para el template
    resumen_prendas = [
        {
            'tipo_prenda': r['tipo_prenda'],
            'talla':       r['talla'],
            'genero':      r['genero'],
            'total':       r['total'],
            'beneficiario': r['pedido__nombre_completo'],
            'cliente':     r['pedido__cliente__nombre'],
            'categoria':   r['pedido__categoria'],
            'gestion':     r['pedido__gestion'],
        }
        for r in items_qs
    ]
 
    # ── Opciones para los selects de filtro ──────────────────────────────
    # Tipos de prenda — ajusta las choices a las de tu modelo
    tipos_prenda = ItemPedidoPersonalizado._meta.get_field('tipo_prenda').choices
 
    # Tallas disponibles (dinámico, solo las que existen)
    tallas_disponibles = (
        ItemPedidoPersonalizado.objects
        .values_list('talla', flat=True)
        .distinct()
        .order_by('talla')
    )
 
    # Géneros — ajusta las choices a las de tu modelo
    generos = ItemPedidoPersonalizado._meta.get_field('genero').choices
 
    return render(request, 'pedidos_personalizados/dashboard.html', {
        # Stats
        'total_personas':    total_personas,
        'total_prendas':     total_prendas,
        'total_aportes':     total_aportes,
        'total_saldos':      total_saldos,
        # Tabla
        'resumen_prendas':   resumen_prendas,
        # Selects
        'tipos_prenda':      tipos_prenda,
        'tallas_disponibles': tallas_disponibles,
        'generos':           generos,
    })
 
# ── 2. NUEVA exportar_word ───────────────────────────────────────────────────
# Requiere: pip install python-docx
@login_required
def exportar_word(request):
    """Exportar pedidos filtrados a un documento Word formateado."""
    queryset = PedidoPersonalizado.objects.all()
    
    # Obtener parámetros de filtro (igual que en la lista)
    tipo_prenda = request.GET.get('tipo_prenda', '').strip()
    talla = request.GET.get('talla', '').strip()
    genero = request.GET.get('genero', '').strip()
    beneficiario = request.GET.get('beneficiario', '').strip()
    cliente = request.GET.get('cliente', '').strip()
    
    # Aplicar filtros usando la relación 'items'
    if tipo_prenda:
        queryset = queryset.filter(items__tipo_prenda=tipo_prenda)
    if talla:
        queryset = queryset.filter(items__talla=talla)
    if genero:
        queryset = queryset.filter(items__genero=genero)
    if beneficiario:
        queryset = queryset.filter(nombre_completo__icontains=beneficiario)
    if cliente:
        queryset = queryset.filter(cliente_id=cliente)
    
    # Eliminar duplicados (un pedido puede tener múltiples items)
    queryset = queryset.distinct()
    
    # Crear documento Word
    doc = Document()
    
    # Título
    title = doc.add_heading('Reporte de Pedidos Personalizados', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo con fecha y filtros aplicados
    from datetime import datetime
    doc.add_paragraph(f'Generado el: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph(f'Total de pedidos: {queryset.count()}')
    
    # Mostrar filtros aplicados (opcional)
    filtros = []
    if tipo_prenda: filtros.append(f'Tipo: {tipo_prenda}')
    if talla: filtros.append(f'Talla: {talla}')
    if genero: filtros.append(f'Género: {genero}')
    if beneficiario: filtros.append(f'Beneficiario: {beneficiario}')
    if cliente: filtros.append(f'Cliente ID: {cliente}')
    if filtros:
        doc.add_paragraph('Filtros aplicados: ' + ', '.join(filtros))
    
    doc.add_paragraph('')
    
    # Tabla de pedidos (7 columnas)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['Cliente', 'Beneficiario', 'Prendas', 'Total Bs.', 'Aporte Bs.', 'Saldo Bs.', 'Estado Pago']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Poner negrita
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    for pedido in queryset:
        # Obtener todas las prendas asociadas a este pedido mediante 'items'
        prendas_qs = pedido.items.all()
        prendas_desc = ', '.join([f"{p.get_tipo_prenda_display()} ({p.talla})" for p in prendas_qs])
        
        row_cells = table.add_row().cells
        row_cells[0].text = pedido.cliente.nombre if pedido.cliente else '—'
        row_cells[1].text = pedido.nombre_completo
        row_cells[2].text = prendas_desc
        row_cells[3].text = f"Bs. {pedido.total or 0:.2f}"
        row_cells[4].text = f"Bs. {pedido.aporte or 0:.2f}"
        row_cells[5].text = f"Bs. {pedido.saldo or 0:.2f}"
        row_cells[6].text = pedido.get_estado_pago_display()
    
    # Preparar respuesta
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="pedidos_personalizados.docx"'
    doc.save(response)
    return response






@login_required
def lista_pedidos(request):
    form_filtros = FiltrosPedidoPersonalizadoForm(request.GET)
    pedidos = PedidoPersonalizado.objects.select_related('cliente').all()

    if form_filtros.is_valid():
        buscar = form_filtros.cleaned_data.get('buscar')
        cliente = form_filtros.cleaned_data.get('cliente')
        gestion = form_filtros.cleaned_data.get('gestion')
        categoria = form_filtros.cleaned_data.get('categoria')
        estado_pedido = form_filtros.cleaned_data.get('estado_pedido')
        estado_pago = form_filtros.cleaned_data.get('estado_pago')

        if buscar:
            pedidos = pedidos.filter(nombre_completo__icontains=buscar)
        if cliente:
            pedidos = pedidos.filter(cliente=cliente)
        if gestion:
            pedidos = pedidos.filter(gestion=gestion)
        if categoria:
            pedidos = pedidos.filter(categoria__icontains=categoria)
        if estado_pedido:
            pedidos = pedidos.filter(estado_pedido=estado_pedido)
        if estado_pago:
            pedidos = pedidos.filter(estado_pago=estado_pago)

    paginator = Paginator(pedidos, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    return render(request, 'pedidos_personalizados/lista.html', {
        'page_obj': page_obj,
        'form_filtros': form_filtros,
    })

@login_required
def crear_pedido(request):
    if request.method == 'POST':
        form = PedidoPersonalizadoForm(request.POST)
        if form.is_valid():
            pedido = form.save()
            messages.success(request, 'Pedido personalizado creado exitosamente.')
            if pedido.tipo_pedido == 'separado':
                return redirect('pedidos_personalizados:agregar_item', pedido_pk=pedido.pk)
            return redirect('pedidos_personalizados:lista')
    else:
        form = PedidoPersonalizadoForm()
    
    return render(request, 'pedidos_personalizados/formulario.html', {'form': form, 'titulo': 'Nuevo Pedido Personalizado'})

@login_required
def editar_pedido(request, pk):
    pedido = get_object_or_404(PedidoPersonalizado, pk=pk)
    if request.method == 'POST':
        form = PedidoPersonalizadoForm(request.POST, instance=pedido)
        if form.is_valid():
            form.save()
            messages.success(request, 'Pedido actualizado.')
            return redirect('pedidos_personalizados:lista')
    else:
        form = PedidoPersonalizadoForm(instance=pedido)
        
    return render(request, 'pedidos_personalizados/formulario.html', {
        'form': form, 
        'titulo': f'Editar Pedido: {pedido.nombre_completo}',
        'pedido': pedido
    })

@login_required
def eliminar_pedido(request, pk):
    pedido = get_object_or_404(PedidoPersonalizado, pk=pk)
    if request.method == 'POST':
        pedido.delete()
        messages.success(request, 'Pedido eliminado.')
        return redirect('pedidos_personalizados:lista')
    return render(request, 'pedidos_personalizados/confirmar_eliminar.html', {'pedido': pedido})

@login_required
def agregar_item(request, pedido_pk):
    pedido = get_object_or_404(PedidoPersonalizado, pk=pedido_pk)
    items = pedido.items.all()
    
    if request.method == 'POST':
        form = ItemPedidoPersonalizadoForm(request.POST)
        if form.is_valid():
            item = form.save(commit=False)
            item.pedido = pedido
            item.save()
            messages.success(request, 'Ítem agregado.')
            return redirect('pedidos_personalizados:agregar_item', pedido_pk=pedido.pk)
    else:
        form = ItemPedidoPersonalizadoForm()
        
    return render(request, 'pedidos_personalizados/agregar_items.html', {
        'pedido': pedido, 'form': form, 'items': items
    })

@login_required
def eliminar_item(request, item_pk):
    item = get_object_or_404(ItemPedidoPersonalizado, pk=item_pk)
    pedido_pk = item.pedido.pk
    item.delete()
    messages.success(request, 'Ítem eliminado.')
    return redirect('pedidos_personalizados:agregar_item', pedido_pk=pedido_pk)

@login_required
def importar_excel(request):
    if request.method == 'POST':
        form = ImportarExcelPersonalizadoForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                df = pd.read_excel(request.FILES['archivo'])
                
                required_cols = ['Cliente', 'Nombre Completo', 'Gestión', 'Categoría', 'Tipo Pedido']
                missing_cols = [c for c in required_cols if c not in df.columns]
                if missing_cols:
                    messages.error(request, f"Faltan columnas requeridas: {', '.join(missing_cols)}")
                    return redirect('pedidos_personalizados:importar')
                
                # Agrupamos por persona
                for index, row in df.iterrows():
                    cliente_nombre = str(row['Cliente']).strip()
                    nombre = str(row['Nombre Completo']).strip()
                    gestion = int(row['Gestión']) if pd.notna(row['Gestión']) else 2024
                    categoria = str(row['Categoría']).strip()
                    tipo_pedido = str(row['Tipo Pedido']).strip().lower()
                    
                    cliente, _ = Cliente.objects.get_or_create(nombre=cliente_nombre)
                    
                    # Tipo pedido puede ser 'conjunto' o 'separado'
                    is_conjunto = 'conjunto' in tipo_pedido
                    t_pedido = 'conjunto' if is_conjunto else 'separado'
                    
                    # Manejar valores numéricos que puedan venir vacíos (NaN)
                    aporte = Decimal(str(row['Aporte'])) if 'Aporte' in df.columns and pd.notna(row['Aporte']) else Decimal('0')
                    total = Decimal(str(row['Total'])) if 'Total' in df.columns and pd.notna(row['Total']) else Decimal('0')
                    
                    # Buscar si ya existe el pedido para agrupar (en caso de prendas separadas en múltiples filas)
                    pedido, created = PedidoPersonalizado.objects.get_or_create(
                        cliente=cliente,
                        nombre_completo=nombre,
                        gestion=gestion,
                        defaults={
                            'categoria': categoria,
                            'tipo_pedido': t_pedido,
                            'aporte': aporte,
                            'total': total,
                            'estado_pedido': 'registrado'
                        }
                    )
                    
                    # Si creó el pedido y era conjunto o separado, se actualiza
                    if created and is_conjunto:
                        talla = str(row['Talla']).strip() if 'Talla' in df.columns and pd.notna(row['Talla']) else 'M'
                        genero = str(row['Género']).strip().lower() if 'Género' in df.columns and pd.notna(row['Género']) else 'unisex'
                        pedido.talla_conjunto = talla
                        pedido.genero_conjunto = genero
                        pedido.save() # auto-generates items
                        
                    elif not is_conjunto:
                        # Si es separado o agrupando
                        if 'Prenda' in df.columns and pd.notna(row['Prenda']):
                            prenda = str(row['Prenda']).strip().lower()
                            talla = str(row['Talla']).strip() if 'Talla' in df.columns and pd.notna(row['Talla']) else 'M'
                            genero = str(row['Género']).strip().lower() if 'Género' in df.columns and pd.notna(row['Género']) else 'unisex'
                            cantidad = int(row['Cantidad']) if 'Cantidad' in df.columns and pd.notna(row['Cantidad']) else 1
                            
                            ItemPedidoPersonalizado.objects.create(
                                pedido=pedido,
                                tipo_prenda=prenda,
                                talla=talla,
                                genero=genero,
                                cantidad=cantidad
                            )
                            
                messages.success(request, 'Archivo procesado exitosamente.')
                return redirect('pedidos_personalizados:lista')
            except Exception as e:
                messages.error(request, f'Error al procesar archivo: {str(e)}')
    else:
        form = ImportarExcelPersonalizadoForm()
        
    return render(request, 'pedidos_personalizados/importar.html', {'form': form})

@login_required
def exportar_excel(request):
    import io
    from django.http import HttpResponse

    form_filtros = FiltrosPedidoPersonalizadoForm(request.GET)
    pedidos = PedidoPersonalizado.objects.select_related('cliente').all()

    if form_filtros.is_valid():
        buscar = form_filtros.cleaned_data.get('buscar')
        cliente = form_filtros.cleaned_data.get('cliente')
        gestion = form_filtros.cleaned_data.get('gestion')
        categoria = form_filtros.cleaned_data.get('categoria')
        estado_pedido = form_filtros.cleaned_data.get('estado_pedido')
        estado_pago = form_filtros.cleaned_data.get('estado_pago')

        if buscar:
            pedidos = pedidos.filter(nombre_completo__icontains=buscar)
        if cliente:
            pedidos = pedidos.filter(cliente=cliente)
        if gestion:
            pedidos = pedidos.filter(gestion=gestion)
        if categoria:
            pedidos = pedidos.filter(categoria__icontains=categoria)
        if estado_pedido:
            pedidos = pedidos.filter(estado_pedido=estado_pedido)
        if estado_pago:
            pedidos = pedidos.filter(estado_pago=estado_pago)

    data = []
    for p in pedidos:
        base_dict = {
            'Cliente': p.cliente.nombre,
            'Nombre Completo': p.nombre_completo,
            'Gestión': p.gestion,
            'Categoría': p.categoria,
            'Tipo Pedido': p.get_tipo_pedido_display(),
            'Aporte': float(p.aporte) if p.aporte else 0,
            'Saldo': float(p.saldo) if p.saldo else 0,
            'Total': float(p.total) if p.total else 0,
            'Estado Pago': p.get_estado_pago_display(),
            'Estado Pedido': p.get_estado_pedido_display(),
        }
        items = p.items.all()
        if not items:
            # Caso sin items
            base_dict['Prenda'] = ''
            base_dict['Talla'] = ''
            base_dict['Género'] = ''
            base_dict['Cantidad'] = 0
            data.append(base_dict)
        else:
            for item in items:
                row = base_dict.copy()
                row['Prenda'] = item.get_tipo_prenda_display()
                row['Talla'] = item.talla
                row['Género'] = item.get_genero_display()
                row['Cantidad'] = item.cantidad
                data.append(row)

    df = pd.DataFrame(data)
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Pedidos Custom')
    
    output.seek(0)
    response = HttpResponse(
        output.read(), 
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename="pedidos_personalizados.xlsx"'
    return response




@login_required
def exportar_pdf(request):
    """
    Exporta el resumen de prendas (con beneficiario y cliente) a PDF.
    Aplica los mismos filtros que el dashboard:
    tipo_prenda, talla, genero, beneficiario, cliente.
 
    Requiere: pip install reportlab
    Agregar a requirements.txt: reportlab>=4.0.0
    """
    import io
    from django.http import HttpResponse
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph,
        Spacer, HRFlowable
    )
    from reportlab.platypus import KeepTogether
    from datetime import datetime
 
    # ── Filtros (mismos parámetros GET que el dashboard) ──────────────────
    tipo_prenda_f  = request.GET.get('tipo_prenda', '').strip()
    talla_f        = request.GET.get('talla', '').strip()
    genero_f       = request.GET.get('genero', '').strip()
    beneficiario_f = request.GET.get('beneficiario', '').strip()
    cliente_f      = request.GET.get('cliente', '').strip()
 
    items_qs = (
        ItemPedidoPersonalizado.objects
        .select_related('pedido', 'pedido__cliente')
        .values(
            'tipo_prenda',
            'talla',
            'genero',
            'pedido__nombre_completo',
            'pedido__cliente__nombre',
            'pedido__categoria',
            'pedido__gestion',
        )
        .annotate(total=Sum('cantidad'))
        .order_by('tipo_prenda', 'talla', 'genero', 'pedido__nombre_completo')
    )
 
    if tipo_prenda_f:
        items_qs = items_qs.filter(tipo_prenda=tipo_prenda_f)
    if talla_f:
        items_qs = items_qs.filter(talla=talla_f)
    if genero_f:
        items_qs = items_qs.filter(genero=genero_f)
    if beneficiario_f:
        items_qs = items_qs.filter(pedido__nombre_completo__icontains=beneficiario_f)
    if cliente_f:
        items_qs = items_qs.filter(pedido__cliente__nombre__icontains=cliente_f)
 
    # Stats globales
    todos         = PedidoPersonalizado.objects.all()
    total_personas = todos.count()
    total_aportes  = todos.aggregate(t=Sum('aporte'))['t'] or 0
    total_saldos   = todos.aggregate(t=Sum('saldo'))['t'] or 0
    total_prendas_filtradas = sum(r['total'] for r in items_qs)
 
    # ── Paleta de colores (misma del sistema dark) ─────────────────────────
    COLOR_BG_HEADER  = colors.HexColor('#161920')   # --surface
    COLOR_BG_ALT     = colors.HexColor('#1e222d')   # --surface2
    COLOR_ACCENT     = colors.HexColor('#f97316')   # naranja
    COLOR_TEXT_MAIN  = colors.HexColor('#f0f2f5')   # --text-1
    COLOR_TEXT_MUTED = colors.HexColor('#8b91a0')   # --text-2
    COLOR_BORDER     = colors.HexColor('#2a2f3d')
    COLOR_WHITE      = colors.white
 
    # ── Documento ─────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=1.5*cm, rightMargin=1.5*cm,
        topMargin=1.5*cm,  bottomMargin=1.5*cm,
        title='Reporte de Pedidos Personalizados',
        author='Sistema GestPed',
    )
 
    styles = getSampleStyleSheet()
 
    style_title = ParagraphStyle(
        'CustomTitle',
        fontName='Helvetica-Bold',
        fontSize=18,
        textColor=COLOR_BG_HEADER,
        spaceAfter=4,
        alignment=TA_LEFT,
    )
 
    style_subtitle = ParagraphStyle(
        'CustomSub',
        fontName='Helvetica',
        fontSize=9,
        textColor=COLOR_TEXT_MUTED,
        spaceAfter=2,
        alignment=TA_LEFT,
    )
 
    style_section = ParagraphStyle(
        'Section',
        fontName='Helvetica-Bold',
        fontSize=11,
        textColor=COLOR_BG_HEADER,
        spaceBefore=14,
        spaceAfter=6,
    )
 
    style_filter_label = ParagraphStyle(
        'FilterLabel',
        fontName='Helvetica',
        fontSize=8,
        textColor=COLOR_TEXT_MUTED,
    )
 
    story = []
 
    # ── Encabezado ─────────────────────────────────────────────────────────
    # Banda de color superior usando una tabla de 1 fila
    header_data = [[
        Paragraph('<b><font color="white" size="16">GestPed</font></b> &nbsp; '
                  '<font color="#f97316" size="9">Reporte de Pedidos Personalizados</font>', styles['Normal']),
        Paragraph(f'<font color="#8b91a0" size="8">{datetime.now().strftime("%d/%m/%Y %H:%M")}</font>',
                  ParagraphStyle('R', alignment=2, fontName='Helvetica', fontSize=8,
                                 textColor=COLOR_TEXT_MUTED)),
    ]]
 
    header_table = Table(header_data, colWidths=['*', 5*cm])
    header_table.setStyle(TableStyle([
        ('BACKGROUND',  (0,0), (-1,-1), COLOR_BG_HEADER),
        ('VALIGN',      (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING',  (0,0), (-1,-1), 10),
        ('BOTTOMPADDING',(0,0),(-1,-1), 10),
        ('LEFTPADDING', (0,0), (-1,-1), 14),
        ('RIGHTPADDING',(-1,0),(-1,-1), 14),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 10))
 
    # ── Tarjetas de stats ──────────────────────────────────────────────────
    stats_data = [[
        Paragraph(f'<b><font size="20">{total_personas}</font></b><br/>'
                  f'<font size="8" color="#8b91a0">Beneficiarios</font>', styles['Normal']),
        Paragraph(f'<b><font size="20">{total_prendas_filtradas}</font></b><br/>'
                  f'<font size="8" color="#8b91a0">Prendas (filtradas)</font>', styles['Normal']),
        Paragraph(f'<b><font size="20">Bs. {total_aportes}</font></b><br/>'
                  f'<font size="8" color="#8b91a0">Total aportes</font>', styles['Normal']),
        Paragraph(f'<b><font size="20" color="#ef4444">Bs. {total_saldos}</font></b><br/>'
                  f'<font size="8" color="#8b91a0">Saldos pendientes</font>', styles['Normal']),
    ]]
 
    stats_table = Table(stats_data, colWidths=['*', '*', '*', '*'])
    stats_table.setStyle(TableStyle([
        ('BACKGROUND',   (0,0), (-1,-1), COLOR_BG_ALT),
        ('BOX',          (0,0), (-1,-1), 0.5, COLOR_BORDER),
        ('INNERGRID',    (0,0), (-1,-1), 0.5, COLOR_BORDER),
        ('VALIGN',       (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING',   (0,0), (-1,-1), 10),
        ('BOTTOMPADDING',(0,0), (-1,-1), 10),
        ('LEFTPADDING',  (0,0), (-1,-1), 14),
        # Línea naranja izquierda solo en primera celda
        ('LINEAFTER',    (0,0), (0,0),   2, COLOR_ACCENT),
    ]))
    story.append(stats_table)
    story.append(Spacer(1, 12))
 
    # ── Filtros activos ────────────────────────────────────────────────────
    filtros_txt = []
    if tipo_prenda_f:  filtros_txt.append(f'Prenda: {tipo_prenda_f}')
    if talla_f:        filtros_txt.append(f'Talla: {talla_f}')
    if genero_f:       filtros_txt.append(f'Género: {genero_f}')
    if beneficiario_f: filtros_txt.append(f'Beneficiario: {beneficiario_f}')
    if cliente_f:      filtros_txt.append(f'Cliente: {cliente_f}')
 
    if filtros_txt:
        story.append(Paragraph(
            '<b>Filtros aplicados:</b> ' + '  ·  '.join(filtros_txt),
            style_filter_label
        ))
        story.append(Spacer(1, 6))
 
    # ── Tabla principal ────────────────────────────────────────────────────
    story.append(Paragraph('Resumen de prendas para producción', style_section))
    story.append(HRFlowable(width='100%', thickness=1, color=COLOR_ACCENT, spaceAfter=8))
 
    # Cabecera de la tabla
    col_headers = ['Prenda', 'Talla', 'Género', 'Cant.', 'Beneficiario', 'Institución', 'Categoría', 'Gestión']
 
    table_data = [col_headers]
 
    for r in items_qs:
        table_data.append([
            r['tipo_prenda'].title(),
            r['talla'],
            r['genero'].title(),
            str(r['total']),
            r['pedido__nombre_completo'],
            r['pedido__cliente__nombre'],
            r['pedido__categoria'] or '—',
            str(r['pedido__gestion']),
        ])
 
    if len(table_data) == 1:
        table_data.append(['Sin resultados', '', '', '', '', '', '', ''])
 
    # Anchos de columna (landscape A4 ≈ 27 cm útiles)
    col_widths = [3*cm, 1.6*cm, 2*cm, 1.5*cm, 5.5*cm, 5.5*cm, 3.5*cm, 2*cm]
 
    main_table = Table(table_data, colWidths=col_widths, repeatRows=1)
 
    # Estilos de la tabla
    ts = TableStyle([
        # Cabecera
        ('BACKGROUND',    (0,0),  (-1,0),  COLOR_BG_HEADER),
        ('TEXTCOLOR',     (0,0),  (-1,0),  COLOR_TEXT_MAIN),
        ('FONTNAME',      (0,0),  (-1,0),  'Helvetica-Bold'),
        ('FONTSIZE',      (0,0),  (-1,0),  8),
        ('TOPPADDING',    (0,0),  (-1,0),  8),
        ('BOTTOMPADDING', (0,0),  (-1,0),  8),
        ('LEFTPADDING',   (0,0),  (-1,0),  8),
        ('ALIGN',         (3,0),  (3,0),   'CENTER'),  # Cant. centrado
        # Línea inferior de cabecera en naranja
        ('LINEBELOW',     (0,0),  (-1,0),  2, COLOR_ACCENT),
 
        # Filas de datos
        ('FONTNAME',      (0,1),  (-1,-1), 'Helvetica'),
        ('FONTSIZE',      (0,1),  (-1,-1), 8),
        ('TOPPADDING',    (0,1),  (-1,-1), 6),
        ('BOTTOMPADDING', (0,1),  (-1,-1), 6),
        ('LEFTPADDING',   (0,1),  (-1,-1), 8),
        ('VALIGN',        (0,0),  (-1,-1), 'MIDDLE'),
        ('TEXTCOLOR',     (0,1),  (-1,-1), COLOR_BG_HEADER),
        ('ALIGN',         (3,1),  (3,-1),  'CENTER'),  # Cant. centrado
 
        # Cantidad en naranja y bold
        ('TEXTCOLOR',     (3,1),  (3,-1),  COLOR_ACCENT),
        ('FONTNAME',      (3,1),  (3,-1),  'Helvetica-Bold'),
        ('FONTSIZE',      (3,1),  (3,-1),  10),
 
        # Grid
        ('LINEBELOW',     (0,1),  (-1,-2), 0.5, COLOR_BORDER),
        ('LINEBELOW',     (0,-1), (-1,-1), 0.5, COLOR_BORDER),
        ('BOX',           (0,0),  (-1,-1), 0.5, COLOR_BORDER),
    ])
 
    # Filas alternas
    for i in range(1, len(table_data)):
        if i % 2 == 0:
            ts.add('BACKGROUND', (0,i), (-1,i), colors.HexColor('#f8f9fa'))
        else:
            ts.add('BACKGROUND', (0,i), (-1,i), colors.white)
 
    main_table.setStyle(ts)
    story.append(main_table)
 
    # ── Pie de página ──────────────────────────────────────────────────────
    story.append(Spacer(1, 16))
    footer_data = [[
        Paragraph(f'<font size="8" color="#8b91a0">Total de registros: <b>{len(table_data)-1}</b></font>',
                  styles['Normal']),
        Paragraph(f'<font size="8" color="#8b91a0">Sistema GestPed · Ropa deportiva para colegios</font>',
                  ParagraphStyle('FooterR', alignment=2, fontName='Helvetica',
                                 fontSize=8, textColor=COLOR_TEXT_MUTED)),
    ]]
    footer_table = Table(footer_data, colWidths=['*', '*'])
    footer_table.setStyle(TableStyle([
        ('LINEABOVE',    (0,0), (-1,0), 0.5, COLOR_BORDER),
        ('TOPPADDING',   (0,0), (-1,-1), 6),
        ('LEFTPADDING',  (0,0), (0,0),  0),
        ('RIGHTPADDING', (-1,0),(-1,0), 0),
    ]))
    story.append(footer_table)
 
    # ── Build y respuesta ──────────────────────────────────────────────────
    doc.build(story)
    buffer.seek(0)
 
    response = HttpResponse(buffer.read(), content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="pedidos_personalizados.pdf"'
    return response

